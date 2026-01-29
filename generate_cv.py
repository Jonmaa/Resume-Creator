"""
ATS-Optimized CV Generator
==========================
Generates a professional, ATS-friendly CV in Word format (.docx) from a JSON data file.

Usage:
    python generate_cv.py [--input cv_data.json] [--output output.docx]

Author: Open Source Project
License: MIT
"""

import json
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class CVGenerator:
    """Generates ATS-optimized CV documents from structured data."""
    
    # Style constants
    COLORS = {
        'primary': RGBColor(0, 51, 102),      # Dark blue
        'secondary': RGBColor(80, 80, 80),     # Gray
        'muted': RGBColor(100, 100, 100),      # Light gray
    }
    
    FONT_SIZES = {
        'name': Pt(26),
        'title': Pt(13),
        'section': Pt(12),
        'body': Pt(11),
        'small': Pt(10),
    }
    
    # Translations for section headers
    TRANSLATIONS = {
        'en': {
            'professional_summary': 'Professional Summary',
            'technical_skills': 'Technical Skills',
            'professional_experience': 'Professional Experience',
            'education': 'Education',
            'certifications_languages': 'Certifications & Languages',
            'certifications': 'Certifications',
            'languages': 'Languages',
        },
        'es': {
            'professional_summary': 'Resumen Profesional',
            'technical_skills': 'Habilidades Técnicas',
            'professional_experience': 'Experiencia Profesional',
            'education': 'Educación',
            'certifications_languages': 'Certificaciones e Idiomas',
            'certifications': 'Certificaciones',
            'languages': 'Idiomas',
        }
    }
    
    def __init__(self, data: dict, language: str = 'en', photo_path: str = None):
        """Initialize with CV data dictionary, language, and optional photo."""
        self.data = data
        self.lang = language if language in self.TRANSLATIONS else 'en'
        self.t = self.TRANSLATIONS[self.lang]
        self.photo_path = photo_path
        self.doc = Document()
        self._setup_document()
    
    def _setup_document(self):
        """Configure document margins for single-page layout."""
        for section in self.doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
    
    def _add_section_heading(self, text: str):
        """Add a styled section heading with bottom border."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = self.FONT_SIZES['section']
        run.font.color.rgb = self.COLORS['primary']
        
        # Add bottom border
        p_border = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:color'), '003366')
        p_border.append(bottom)
        p._p.get_or_add_pPr().append(p_border)
        
        return p
    
    def _add_bullet(self, text: str):
        """Add a compact bullet point."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.line_spacing = 1.15
        
        run = p.add_run(f"• {text}")
        run.font.size = self.FONT_SIZES['body']
        return p
    
    def _remove_table_borders(self, table):
        """Remove all borders from a table."""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

    def build_header(self):
        """Build the header section with name, title, contact info, and optional photo."""
        personal = self.data.get('personal', {})
        
        # Check if photo exists
        has_photo = self.photo_path and Path(self.photo_path).exists()
        
        if has_photo:
            # Create table with 2 columns: info (left) + photo (right)
            table = self.doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            self._remove_table_borders(table)
            
            # Set column widths
            table.columns[0].width = Inches(5.0)
            table.columns[1].width = Inches(1.5)
            
            left_cell = table.rows[0].cells[0]
            right_cell = table.rows[0].cells[1]
            
            # Left cell: Name, title, contact
            name_p = left_cell.paragraphs[0]
            name_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            name_run = name_p.add_run(personal.get('name', 'Your Name').upper())
            name_run.bold = True
            name_run.font.size = self.FONT_SIZES['name']
            name_run.font.color.rgb = self.COLORS['primary']
            
            title_p = left_cell.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_p.paragraph_format.space_before = Pt(2)
            title_p.paragraph_format.space_after = Pt(6)
            title_run = title_p.add_run(personal.get('title', 'Professional Title'))
            title_run.font.size = self.FONT_SIZES['title']
            title_run.font.color.rgb = self.COLORS['secondary']
            
            # Contact line 1
            contact_parts = []
            if personal.get('email'):
                contact_parts.append(personal['email'])
            if personal.get('phone'):
                contact_parts.append(personal['phone'])
            if personal.get('location'):
                contact_parts.append(personal['location'])
            
            if contact_parts:
                contact_p = left_cell.add_paragraph()
                contact_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                contact_p.paragraph_format.space_after = Pt(2)
                contact_run = contact_p.add_run("  |  ".join(contact_parts))
                contact_run.font.size = self.FONT_SIZES['small']
            
            # Links line 2
            link_parts = []
            if personal.get('github'):
                link_parts.append(personal['github'])
            if personal.get('linkedin'):
                link_parts.append(personal['linkedin'])
            if personal.get('portfolio'):
                link_parts.append(personal['portfolio'])
            
            if link_parts:
                links_p = left_cell.add_paragraph()
                links_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                links_p.paragraph_format.space_before = Pt(0)
                links_run = links_p.add_run("  |  ".join(link_parts))
                links_run.font.size = self.FONT_SIZES['small']
            
            # Right cell: Photo
            photo_p = right_cell.paragraphs[0]
            photo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = photo_p.add_run()
            run.add_picture(self.photo_path, height=Inches(1.4))
            
            # Add spacing after table
            spacing_p = self.doc.add_paragraph()
            spacing_p.paragraph_format.space_after = Pt(6)
        
        else:
            # No photo: centered layout (original)
            name_p = self.doc.add_paragraph()
            name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_p.paragraph_format.space_after = Pt(2)
            name_run = name_p.add_run(personal.get('name', 'Your Name').upper())
            name_run.bold = True
            name_run.font.size = self.FONT_SIZES['name']
            name_run.font.color.rgb = self.COLORS['primary']
            
            title_p = self.doc.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_p.paragraph_format.space_before = Pt(0)
            title_p.paragraph_format.space_after = Pt(6)
            title_run = title_p.add_run(personal.get('title', 'Professional Title'))
            title_run.font.size = self.FONT_SIZES['title']
            title_run.font.color.rgb = self.COLORS['secondary']
            
            # Contact line 1
            contact_parts = []
            if personal.get('email'):
                contact_parts.append(personal['email'])
            if personal.get('phone'):
                contact_parts.append(personal['phone'])
            if personal.get('location'):
                contact_parts.append(personal['location'])
            
            if contact_parts:
                contact_p = self.doc.add_paragraph()
                contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact_p.paragraph_format.space_after = Pt(2)
                contact_run = contact_p.add_run("  |  ".join(contact_parts))
                contact_run.font.size = self.FONT_SIZES['small']
            
            # Links line 2
            link_parts = []
            if personal.get('github'):
                link_parts.append(personal['github'])
            if personal.get('linkedin'):
                link_parts.append(personal['linkedin'])
            if personal.get('portfolio'):
                link_parts.append(personal['portfolio'])
            
            if link_parts:
                links_p = self.doc.add_paragraph()
                links_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                links_p.paragraph_format.space_before = Pt(0)
                links_p.paragraph_format.space_after = Pt(10)
                links_run = links_p.add_run("  |  ".join(link_parts))
                links_run.font.size = self.FONT_SIZES['small']
    
    def build_summary(self):
        """Build the professional summary section."""
        summary_text = self.data.get('summary', '')
        if not summary_text:
            return
        
        self._add_section_heading(self.t['professional_summary'])
        
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(summary_text)
        run.font.size = self.FONT_SIZES['body']
    
    def build_skills(self):
        """Build the technical skills section."""
        skills = self.data.get('skills', {})
        if not skills:
            return
        
        self._add_section_heading(self.t['technical_skills'])
        
        # Compact format: Category: skill1, skill2, skill3
        for category, skill_list in skills.items():
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.1
            
            cat_run = p.add_run(f"{category}: ")
            cat_run.bold = True
            cat_run.font.size = self.FONT_SIZES['body']
            
            skills_run = p.add_run(skill_list)
            skills_run.font.size = self.FONT_SIZES['body']
    
    def build_experience(self):
        """Build the professional experience section."""
        experience = self.data.get('experience', [])
        if not experience:
            return
        
        self._add_section_heading(self.t['professional_experience'])
        
        for job in experience:
            # Job header
            header_p = self.doc.add_paragraph()
            header_p.paragraph_format.space_before = Pt(8)
            header_p.paragraph_format.space_after = Pt(2)
            
            title_run = header_p.add_run(job.get('title', ''))
            title_run.bold = True
            title_run.font.size = self.FONT_SIZES['body']
            
            header_p.add_run(" | ").font.size = self.FONT_SIZES['body']
            
            company_run = header_p.add_run(job.get('company', ''))
            company_run.italic = True
            company_run.font.size = self.FONT_SIZES['body']
            
            header_p.add_run(f" | {job.get('location', '')}").font.size = self.FONT_SIZES['body']
            
            # Date
            date_p = self.doc.add_paragraph()
            date_p.paragraph_format.space_before = Pt(0)
            date_p.paragraph_format.space_after = Pt(4)
            date_run = date_p.add_run(job.get('dates', ''))
            date_run.font.size = self.FONT_SIZES['small']
            date_run.italic = True
            date_run.font.color.rgb = self.COLORS['muted']
            
            # Achievements
            for achievement in job.get('achievements', []):
                self._add_bullet(achievement)
    
    def build_education(self):
        """Build the education section."""
        education = self.data.get('education', [])
        if not education:
            return
        
        self._add_section_heading(self.t['education'])
        
        for edu in education:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            
            degree_run = p.add_run(edu.get('degree', ''))
            degree_run.bold = True
            degree_run.font.size = self.FONT_SIZES['body']
            
            details_p = self.doc.add_paragraph()
            details_p.paragraph_format.space_after = Pt(3)
            details_run = details_p.add_run(f"{edu.get('institution', '')} | {edu.get('dates', '')}")
            details_run.font.size = self.FONT_SIZES['body']
            details_run.italic = True
            
            if edu.get('details'):
                self._add_bullet(edu['details'])
    
    def build_certifications_and_languages(self):
        """Build certifications and languages in a compact combined section."""
        certs = self.data.get('certifications', [])
        languages = self.data.get('languages', [])
        
        if not certs and not languages:
            return
        
        # Combined section header
        self._add_section_heading(self.t['certifications_languages'])
        
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        
        parts = []
        
        # Certifications inline
        if certs:
            parts.append(f"{self.t['certifications']}: {', '.join(certs)}")
        
        # Languages inline
        if languages:
            lang_strs = [f"{l['language']} ({l['level']})" for l in languages]
            parts.append(f"{self.t['languages']}: {', '.join(lang_strs)}")
        
        run = p.add_run("  •  ".join(parts))
        run.font.size = self.FONT_SIZES['body']
    
    def generate(self, output_path: str):
        """Generate the complete CV document."""
        self.build_header()
        self.build_summary()
        self.build_skills()
        self.build_experience()
        self.build_education()
        self.build_certifications_and_languages()
        
        self.doc.save(output_path)
        return output_path


def main():
    parser = argparse.ArgumentParser(
        description='Generate an ATS-optimized CV in Word format from JSON data.',
        epilog='Example: python generate_cv.py --input my_cv.json --output my_cv.docx'
    )
    parser.add_argument(
        '--input', '-i',
        default='cv_data.json',
        help='Path to the JSON file containing CV data (default: cv_data.json)'
    )
    parser.add_argument(
        '--output', '-o',
        default='CV_Optimized_ATS.docx',
        help='Output path for the generated Word document (default: CV_Optimized_ATS.docx)'
    )
    parser.add_argument(
        '--lang', '-l',
        default='en',
        choices=['en', 'es'],
        help='Language for section headers: en (English) or es (Spanish) (default: en)'
    )
    parser.add_argument(
        '--photo', '-p',
        default=None,
        help='Path to a profile photo (jpg/png) to include in the header (optional)'
    )
    
    args = parser.parse_args()
    
    # Resolve paths
    script_dir = Path(__file__).parent
    input_path = Path(args.input)
    output_path = Path(args.output)
    photo_path = None
    
    if not input_path.is_absolute():
        input_path = script_dir / input_path
    if not output_path.is_absolute():
        output_path = script_dir / output_path
    if args.photo:
        photo_path = Path(args.photo)
        if not photo_path.is_absolute():
            photo_path = script_dir / photo_path
    
    # Load data
    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}")
        print("   Create a cv_data.json file with your CV data or specify --input path")
        return 1
    
    print(f"Loading CV data from: {input_path.name}")
    
    with open(input_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Generate CV
    lang_name = 'English' if args.lang == 'en' else 'Spanish'
    photo_msg = f" with photo" if photo_path and photo_path.exists() else ""
    print(f"Generating ATS-optimized CV in {lang_name}{photo_msg}...")
    generator = CVGenerator(data, language=args.lang, photo_path=str(photo_path) if photo_path else None)
    generator.generate(str(output_path))
    
    print(f"CV successfully generated: {output_path.name}")
    print(f"Full path: {output_path}")
    
    return 0


if __name__ == "__main__":
    exit(main())
